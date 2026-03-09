from __future__ import annotations

import io

from docx import Document


def markdown_to_docx(markdown_text: str) -> bytes:
    doc = Document()
    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
            continue
        doc.add_paragraph(line)

    buff = io.BytesIO()
    doc.save(buff)
    return buff.getvalue()
